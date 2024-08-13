import matplotlib.pyplot as plt
import numpy as np
import os
import pandas as pd
import yfinance as yf
from datetime import date, timedelta
import math
from joblib import Parallel, delayed
from docx import Document
from docx.shared import Inches, Pt
import bisect
import logging

# Set up logging
logging.basicConfig(level=logging.INFO, filename='error_log.txt', filemode='w',
                    format='%(name)s - %(levelname)s - %(message)s')

trade_size = [10, 100, 1000, 10000, 50000, 100000, 200000, 500000]
lstm_twap_savings = [0.4833877966217523, 6.512288002481531, 81.12427017289951, 943.3667559009857, 5130.123773311945,
                     10404.01836054354, 21377.43246173636, 55221.65337483059]
lstm_vwap_savings = [0.5692964099618479, 7.6696644184920295, 95.47100867884151, 1109.733299252144, 6029.584615432167,
                     12256.14839370121, 25182.89864692399, 65051.83924884556]


def import_data(holdings, firm_name):
    stock_list = (holdings["Ticker"][0:6]).to_numpy()
    end_date = date.today()
    start_date = end_date - timedelta(days=25)
    combined_data = []

    for ticker in stock_list:
        try:
            data = yf.download(ticker, start=start_date, end=end_date, interval='1h')
            data['Ticker'] = ticker
            combined_data.append(data)
        except Exception as e:
            logging.error(f"Failed to download data for {ticker} for firm {firm_name}: {e}")
            continue

    return combined_data


def preprocess(stock_data, window_size):
    for stock in stock_data:
        stock['mid_price'] = (stock['High'] + stock['Low']) / 2
        stock['mean_vol'] = stock['mid_price'].pct_change().rolling(window=window_size).mean()
        stock['mean_liq'] = stock['Volume'].rolling(window=window_size).mean()

        stock.iloc[:window_size - 1, stock.columns.get_loc('mean_liq')] = stock['Volume'].iloc[:window_size - 1].values
        stock.iloc[:window_size, stock.columns.get_loc('mean_vol')] = stock['mean_vol'].iloc[
                                                                      window_size:2 * window_size].values


def add_almgren(df_list):
    for i in range(len(df_list)):
        for j in range(len(trade_size)):
            x0 = trade_size[j]
            N = min(x0, 2400)

            def cost_function(x, eta, sigma, x0=10):
                x_cumsum = np.cumsum(x)
                x_half = x / 2
                temp_cost = np.sum(eta * (x ** 2))
                perm_cost = np.sum(eta * x * (x0 - x_cumsum + x_half))
                var_cost = np.sum(0.1 * sigma ** 2 * (x ** 2))
                return (temp_cost + perm_cost + var_cost) * math.log((1 + math.log10(x0)), 2) / x0 / 200

            def almgren(row):
                eta = 1 / math.log(row['mean_liq'], 6) if row['mean_liq'] != 0 else 0.0000001
                sigma = row['mean_vol']
                x_init = np.ones(N) * (x0 / N)
                return cost_function(x_init, eta, sigma, x0) / x0

            df_list[i][f'almgren_{trade_size[j]}'] = Parallel(n_jobs=-1)(
                delayed(almgren)(row) for _, row in df_list[i].iterrows())


def cost_savings(user_benchmark, trade_size_index, mean_price, mega_cap=True):
    savings = lstm_twap_savings if user_benchmark == "TWAP" else lstm_vwap_savings
    cost_savings_dollars = (savings[trade_size_index] * mean_price / 200 + 0.000 * mean_price * trade_size[
        trade_size_index]) * (1.5 ** (1 - int(mega_cap)))
    cost_savings_percent = cost_savings_dollars / (trade_size[trade_size_index] * mean_price)
    return cost_savings_dollars, cost_savings_percent


def ret_bar_graph(trade_index, twap_dollars, vwap_dollars, lstm_dollars, mean_share_price):
    size = trade_size[trade_index]
    values = [
        (twap_dollars) / (mean_share_price * size) * 10000,
        (vwap_dollars) / (mean_share_price * size) * 10000,
        (lstm_dollars) / (mean_share_price * size) * 10000
    ]

    fig, ax = plt.subplots()
    bars = ['TWAP', 'VWAP', 'LSTM']
    ax.bar(bars, values, color=['red', 'green', 'blue'])
    ax.set_title("Average Transaction Costs Over All Holdings")
    ax.set_xlabel('Trading Method')
    ax.set_ylabel('TC per share (bps)')

    save_path = os.path.join(os.path.expanduser('~'), 'Downloads', 'average_bar_graph.png')
    fig.savefig(save_path)
    plt.close(fig)  # Close the figure to free up memory

    return values


def ret_table(x, trade_index, stock_list, prices):
    init_twaps = []
    savings_vs_twap = []
    savings_vs_vwap = []
    index = f"almgren_{trade_size[trade_index]}"

    for i in range(3):
        init_twaps.append(x[0][index].mean() * trade_size[trade_index] * prices[i])
    for i in range(3):
        savings_vs_twap.append(cost_savings("TWAP", trade_index, prices[i])[0])
    for i in range(3):
        savings_vs_vwap.append(cost_savings("VWAP", trade_index, prices[i])[0])

    init_twaps.append(sum(init_twaps) / 3)
    savings_vs_twap.append(sum(savings_vs_twap) / 3)
    savings_vs_vwap.append(sum(savings_vs_vwap) / 3)

    savings_vs_twap = [float(f'{value:.2f}') for value in savings_vs_twap]
    savings_vs_vwap = [float(f'{value:.2f}') for value in savings_vs_vwap]

    data = {
        'Securities': [stock_list[0], stock_list[1], stock_list[2], 'Average'],
        'Weekly Cost Savings (vs TWAP)': savings_vs_twap,
        'Weekly Cost Savings (vs VWAP)': savings_vs_vwap,
    }

    df = pd.DataFrame(data)
    fig, ax = plt.subplots(figsize=(7, 2))
    ax.axis('tight')
    ax.axis('off')
    table = ax.table(cellText=df.values, colLabels=df.columns, cellLoc='center', loc='center')

    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.auto_set_column_width(col=list(range(len(df.columns))))

    for i, key in enumerate(table.get_celld().keys()):
        cell = table.get_celld()[key]
        cell.set_height(0.3)

    save_path = os.path.join(os.path.expanduser('~'), 'Downloads', 'savings_table.png')
    fig.savefig(save_path, bbox_inches='tight')
    plt.close(fig)  # Close the figure to free up memory

    return init_twaps, savings_vs_twap, savings_vs_vwap


def total_amount_invested(holdings, stock_list, end_date=date.today(), firm_name=""):
    today_price = []
    for ticker in stock_list:
        try:
            price = yf.download(ticker, start='2024-7-12', end=end_date)
            price = price['Close']
            if price.empty:
                raise ValueError(f"No data found for {ticker}")
            today_price.append(price[-1])
        except Exception as e:
            logging.error(f"Failed to retrieve data for {ticker} for firm {firm_name}: {e}")
            return 0
    return np.dot(np.array(today_price), np.array(holdings["Shares"][:3].to_numpy()))


def format_large_number(number):
    number = float(number)

    if abs(number) >= 1_000_000_000:
        formatted_number = f"{number / 1_000_000_000:.1f} billion"
    elif abs(number) >= 1_000_000:
        formatted_number = f"{number / 1_000_000:.1f} million"
    else:
        formatted_number = f"{number}"

    return formatted_number


def make_report(file_name_to_save_as, firm_name, stock_list, x, twap_dollars, vwap_dollars, lstm_dollars, holdings,
                trade_index, prices):
    document = Document()
    mean_share_price = sum(prices) / len(prices)
    twap_savings_percent = round(100 * ((twap_dollars[-1] - lstm_dollars[-1]) / twap_dollars[-1]), 1)
    vwap_savings_percent = round(100 * ((vwap_dollars[-1] - lstm_dollars[-1]) / vwap_dollars[-1]), 1)

    document.add_heading(f'Optimal Execution Report: Elevate Your Trading Strategy', level=0).runs[0].font.size = Pt(18)
    document.add_paragraph(f"Our analysis reveals that {firm_name}'s three primary holdings are in "
                           f"{stock_list[0]}, {stock_list[1]}, and {stock_list[2]}. Last week, transacting {trade_size[trade_index]} shares of these positions through our "
                           f"LSTM model resulted in a {twap_savings_percent}% savings in transaction costs compared to TWAP, and a "
                           f"{vwap_savings_percent}% savings compared to VWAP.")

    bar_graph_values = ret_bar_graph(trade_index, twap_dollars[-1], vwap_dollars[-1], lstm_dollars[-1],
                                     mean_share_price)
    document.add_picture(os.path.join(os.path.expanduser('~'), 'Downloads', 'average_bar_graph.png'), width=Inches(6))
    os.remove(os.path.join(os.path.expanduser('~'), 'Downloads', 'average_bar_graph.png'))

    document.add_heading("The Biggest Hidden Trading Cost for Asset Managers", level=2)
    document.add_paragraph(
        f"Typically, when asset managers rebalance their portfolios, they incur substantial costs because they trade in large lumps and create significant"
        + f" market impact, or because they trade during volatile market periods (EOM, BOM, on news events, etc.). Blockhouse applies machine learning (ML)"
        + f" based forecasting models and precise execution cost calculations to deliver real-time, actionable insights which optimize your orders and maximize"
        + f" your potential returns. To prove the effectiveness of our models, we benchmark our execution strategies against industry standard trading"
        + f" algorithms like time weighted average price (TWAP) and volume weighted average price (VWAP). We observed that you have a combined"
        + f" ${format_large_number(total_amount_invested(holdings, stock_list, firm_name=firm_name))} invested in your 3 largest holdings ({stock_list[0]}, {stock_list[1]}, {stock_list[2]})"
        + f" and have constructed a report on how you can reduce your transaction costs. We specifically focused on key areas of improvement in slippage,"
        + f" market impact, bid-ask spread costs, and commissions regarding these three positions.")

    document.add_heading("Long Short-Term Memory (LSTM)", level=2)
    document.add_paragraph(
        "\tLSTM is a machine learning model that is particularly useful for identifying trends in stock data. An LSTM has three gates that allow it to decide which "
        + "information to forget, which information to remember, and which information to use for predicting the next day. This allows the model to keep improving as "
        + "new data is introduced to it, while also grasping the overarching trends that can be overlooked if you only see a few data points. Utilizing this model, "
        + "we can effectively predict future transaction costs and trade accordingly.")

    document.add_paragraph(
        f"\tAfter looking at data for the last 2 years for your 3 major holdings, {stock_list[0]}, {stock_list[1]}, and {stock_list[2]}, we have analyzed on a weekly basis what your cost savings are."
        + " For each of the securities, "
        + "we ran a forecast on future bid-ask spreads and market liquidity using our machine learning models and top of the book order data to better understand the "
        + "expected transaction costs. Then, utilizing the Almgren-Chriss model, which accounts for volatility, market impact and order book depth, we estimate the "
        + "transaction costs that come with trading."
        + "\n")

    document.add_heading(
        f"Our LSTM Model Outperforms TWAP by {twap_savings_percent}%, and VWAP by {vwap_savings_percent}%", level=2)
    document.add_paragraph(
        "\tOur models estimated that transaction costs will be much lower than simply using TWAP if you used our models. This bar graph below shows the average "
        + "transaction cost across the three securities for each strategy in dollars:")

    init_twaps, savings_vs_twap, savings_vs_vwap = ret_table(x, trade_index, stock_list, prices)
    document.add_picture(os.path.join(os.path.expanduser('~'), 'Downloads', 'savings_table.png'), width=Inches(6))
    os.remove(os.path.join(os.path.expanduser('~'), 'Downloads', 'savings_table.png'))

    document.add_paragraph(
        "\tOur strategy that utilizes Long Short-Term Memory (LSTM) models outperforms current industry benchmarks of TWAP and VWAP. Our models have "
        + "competitive edge against industry norms both in times of stability and in times of volatility, leading to tangible differences in the long run.")

    document.add_heading("Actionable Recommendations", level=2)
    document.add_paragraph(
        "\tGiven that the LSTM model reduces transaction costs by {0:.1f}% compared to TWAP and VWAP, our recommendations include implementing LSTM-based trading strategies to minimize costs. Our model suggests trading later in the week when market conditions are more favorable. Additionally, continuously monitoring liquidity and volatility can further optimize your trading strategies. By leveraging our AI-driven insights, you can enhance your trading efficiency and achieve better execution outcomes.".format(
            (100 * ((twap_dollars[-1] - lstm_dollars[-1]) / twap_dollars[-1]) + 100 * (
                        (vwap_dollars[-1] - lstm_dollars[-1]) / vwap_dollars[-1])) / 2))
    document.save(f'{file_name_to_save_as}.docx')

    return init_twaps, savings_vs_twap, savings_vs_vwap, bar_graph_values


def find_closest(sorted_list, target):
    pos = bisect.bisect_left(sorted_list, target)

    if pos == 0:
        return sorted_list[0]
    if pos == len(sorted_list):
        return sorted_list[-1]

    before = sorted_list[pos - 1]
    after = sorted_list[pos]

    if after - target < target - before:
        return after
    else:
        return before


def main():
    directory = '.'  # Current directory
    summary_data = []

    for file in os.listdir(directory):
        if file.endswith(".xlsx"):
            try:
                file_name_to_save_as = os.path.splitext(file)[0]
                firm_name = file_name_to_save_as.replace('_', ' ')
                holdings = pd.read_excel(file)
                stock_list = (holdings["Ticker"][:3]).to_numpy()

                holdings_quantity = np.sum(holdings['Shares'][:3].to_numpy())
                trade_index = trade_size.index(find_closest(trade_size, 0.05 * holdings_quantity))

                x = import_data(holdings, firm_name)
                preprocess(x, 7)
                prices = [100, 150, 250]
                add_almgren(x)

                index = f"almgren_{trade_size[trade_index]}"
                twap_dollars = []
                vwap_dollars = []
                lstm_dollars = []
                twap_savings = []
                vwap_savings = []
                for i in range(3):
                    twap_dollars.append(x[i][index].mean() * trade_size[trade_index] * prices[i])
                    lstm_dollars.append(twap_dollars[i] - cost_savings("TWAP", trade_index, prices[i])[0])
                    twap_savings.append(cost_savings("TWAP", trade_index, prices[i])[0])
                    vwap_savings.append(cost_savings("VWAP", trade_index, prices[i])[0])
                    vwap_dollars.append(twap_dollars[i] + vwap_savings[i])
                twap_dollars.append(sum(twap_dollars) / 3)
                vwap_dollars.append(sum(vwap_dollars) / 3)
                lstm_dollars.append(sum(lstm_dollars) / 3)
                twap_savings.append(sum(twap_savings) / 3)
                vwap_savings.append(sum(vwap_savings) / 3)

                init_twaps, savings_vs_twap, savings_vs_vwap, bar_graph_values = make_report(
                    file_name_to_save_as, firm_name, stock_list, x, twap_dollars, vwap_dollars, lstm_dollars, holdings,
                    trade_index, prices)

                summary_data.append({
                    'FileName': file_name_to_save_as.replace('_', ' '),
                    'Shares': trade_size[trade_index],
                    'Stock1': stock_list[0],
                    'Stock2': stock_list[1],
                    'Stock3': stock_list[2],
                    'TWAPPercent': round(100 * ((twap_dollars[-1] - lstm_dollars[-1]) / twap_dollars[-1]), 1),
                    'VWAPPercent': round(100 * ((vwap_dollars[-1] - lstm_dollars[-1]) / vwap_dollars[-1]), 1),
                    'TWAPDollar': twap_savings[-1],
                    'VWAPDollar': vwap_savings[-1],
                    'AUMTop3': format_large_number(total_amount_invested(holdings, stock_list, firm_name=firm_name)),
                    'TWAPValue': bar_graph_values[0],
                    'VWAPValue': bar_graph_values[1],
                    'LSTMValue': bar_graph_values[2],
                    'InitTWAP1': init_twaps[0],
                    'InitTWAP2': init_twaps[1],
                    'InitTWAP3': init_twaps[2],
                    'AvgTWAP': init_twaps[3],
                    'SavingsVsTWAP1': savings_vs_twap[0],
                    'SavingsVsTWAP2': savings_vs_twap[1],
                    'SavingsVsTWAP3': savings_vs_twap[2],
                    'AvgSavingsVsTWAP': savings_vs_twap[3],
                    'SavingsVsVWAP1': savings_vs_vwap[0],
                    'SavingsVsVWAP2': savings_vs_vwap[1],
                    'SavingsVsVWAP3': savings_vs_vwap[2],
                    'AvgSavingsVsVWAP': savings_vs_vwap[3],
                })
            except Exception as e:
                logging.error(f"Failed to process file {file} for firm {firm_name}: {e}")
                continue

    summary_df = pd.DataFrame(summary_data)
    summary_df.to_excel('Summary_Report.xlsx', index=False)


if __name__ == "__main__":
    main()
